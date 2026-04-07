"""
app.py

This is the entrypoint for the Site Visit Report tool, packaged as a Flask
application suitable for deployment on Render or any other Python web
hosting platform.  It reuses the implementation previously provided in
``site_visit_report_app.py`` and renames the module and variables to
conform to typical deployment conventions.  The app collects site visit
information from Solution Engineers, calls an MCP agent (placeholder),
and generates a Word document report based on the Cloud Inventory site
survey template.

To run this app locally, install the requirements and execute

    python app.py

On Render, the service can be configured to install the requirements
from ``requirements.txt`` and start the web service using Gunicorn with
``gunicorn app:app``.
"""

import os
import uuid
from datetime import datetime

from flask import Flask, render_template_string, request, send_file, redirect, url_for
from docx import Document
from docx.shared import Inches


app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB upload limit
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploaded_photos')
app.config['DATA_FOLDER'] = os.path.join(os.path.dirname(__file__), 'records')

# Ensure upload and data directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DATA_FOLDER'], exist_ok=True)

# Path to a JSON file storing Solution Engineer names for the dropdown.
SE_NAMES_FILE = os.path.join(app.config['DATA_FOLDER'], 'se_names.json')

# Ensure the SE names file exists.
if not os.path.exists(SE_NAMES_FILE):
    # Initialise with an empty list
    with open(SE_NAMES_FILE, 'w', encoding='utf-8') as f:
        f.write('[]')


def call_mcp_agent(summary: str) -> str:
    """Placeholder for calling the Cloud Inventory MCP agent.

    In a real implementation, this function would send the ``summary``
    (e.g., a description of the business area or issue) to your MCP
    agent API and return the agent’s response.  Here we simply return
    a canned response for demonstration purposes.

    Args:
        summary: A text summary of observations and issues.

    Returns:
        A descriptive string explaining how Cloud Inventory can address
        the identified problems.
    """
    return (
        "Based on the information provided, Cloud Inventory can streamline "
        "your operations by automating receiving and putaway via RF "
        "scanners, implementing directed picking and packing workflows, "
        "and offering real-time inventory visibility through dashboards "
        "and cycle counting. This will reduce manual data entry errors and "
        "improve order accuracy."
    )


def generate_word_report(data: dict, image_paths: list[str], mcp_response: str) -> str:
    """Generate a Word document for the site visit report.

    The report includes sections based on the site survey template and
    embeds uploaded images.  It returns the path to the generated file.

    Args:
        data: Dictionary of form data from the site visit submission.
        image_paths: List of file paths to uploaded images.
        mcp_response: The response text from the MCP agent.

    Returns:
        The absolute path to the generated ``.docx`` file.
    """
    filename = f"site_survey_report_{uuid.uuid4().hex}.docx"
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    doc = Document()
    doc.add_heading('Site Survey Report', level=0)
    # Cover page
    doc.add_paragraph(
        f"Survey location: {data.get('survey_location','')}\n"
        f"Survey conducted on: {data.get('survey_date','')}\n"
        f"Company name: {data.get('company_name','')}\n"
        f"Generated on: {datetime.now().strftime('%Y-%m-%d')}"
    )
    doc.add_page_break()
    # Company & Site Profile
    # 1. Company & Site Profile
    doc.add_heading('1. Company & Site Profile', level=1)
    doc.add_paragraph(data.get('company_profile', ''))
    # Add any image associated with the company profile
    if data.get('company_profile_image'):
        try:
            doc.add_picture(os.path.join(app.static_folder, data['company_profile_image']), width=Inches(4))
            doc.add_paragraph('')
        except Exception:
            pass
    # 2. Master Data
    doc.add_heading('2. Master Data', level=1)
    doc.add_paragraph(data.get('master_data_text', ''))
    if data.get('master_data_image'):
        try:
            doc.add_picture(os.path.join(app.static_folder, data['master_data_image']), width=Inches(4))
            doc.add_paragraph('')
        except Exception:
            pass
    # 3. Inventory Information
    doc.add_heading('3. Inventory Information', level=1)
    doc.add_paragraph(data.get('inventory_text', ''))
    if data.get('inventory_image'):
        try:
            doc.add_picture(os.path.join(app.static_folder, data['inventory_image']), width=Inches(4))
            doc.add_paragraph('')
        except Exception:
            pass
    # 4. IT & Systems Landscape
    doc.add_heading('4. IT & Systems Landscape', level=1)
    doc.add_paragraph(data.get('it_systems', ''))
    # 5. Stakeholders (Company and CI representatives)
    doc.add_heading('5. Stakeholders', level=1)
    company_reps = data.get('company_representatives', [])
    ci_reps = data.get('ci_representatives', [])
    if company_reps:
        doc.add_paragraph('Company Representatives:', style=None)
        for rep in company_reps:
            name = rep.get('name', '')
            title = rep.get('title', '')
            doc.add_paragraph(f"- {name} ({title})", style='List Bullet')
    if ci_reps:
        doc.add_paragraph('Cloud Inventory Representatives:', style=None)
        for rep in ci_reps:
            name = rep.get('name', '')
            title = rep.get('title', '')
            doc.add_paragraph(f"- {name} ({title})", style='List Bullet')
    # 6. Operations details
    doc.add_heading('6. Operations', level=1)
    operations = data.get('operations', {})
    for sec_name, sec_data in operations.items():
        doc.add_heading(sec_name, level=2)
        doc.add_paragraph(f"Observations: {sec_data.get('observations','')}")
        doc.add_paragraph(f"Pain point: {sec_data.get('pain','')}")
        doc.add_paragraph(f"Root cause: {sec_data.get('root_cause','')}")
        doc.add_paragraph(f"Business impact: {sec_data.get('impact','')}")
        # Images for this section
        for rel_img in sec_data.get('images', []):
            # Convert relative path to absolute path
            img_path = os.path.join(app.static_folder, rel_img)
            try:
                doc.add_picture(img_path, width=Inches(4))
                doc.add_paragraph('')
            except Exception:
                pass
    # Proposed Solution & Benefits
    doc.add_heading('7. Proposed Solution & Benefits', level=1)
    doc.add_paragraph(mcp_response)
    # In‑Scope & Out‑of‑Scope
    doc.add_heading('8. In‑Scope & Out‑of‑Scope Requirements', level=1)
    doc.add_paragraph(f"In scope: {data.get('in_scope','')}")
    doc.add_paragraph(f"Out of scope: {data.get('out_scope','')}")
    # Value Summary & Metrics
    doc.add_heading('9. Value Summary & Metrics', level=1)
    doc.add_paragraph(data.get('value_summary',''))
    # Attach general images
    if image_paths:
        doc.add_heading('10. Photos', level=1)
        for img_path in image_paths:
            try:
                doc.add_picture(img_path, width=Inches(5))
                doc.add_paragraph('')
            except Exception:
                pass
    # Sign-Off
    doc.add_heading('11. Sign‑Off', level=1)
    doc.add_paragraph('Customer representative signature: ________________')
    doc.add_paragraph('Solution Engineer signature: ________________')
    doc.save(filepath)
    return filepath


# -----------------------------------------------------------------------------
# Data persistence helpers
#
def _record_filepath(record_id: str) -> str:
    """Return the full path to the JSON file for a given record ID."""
    return os.path.join(app.config['DATA_FOLDER'], f"{record_id}.json")


def save_record(record_id: str, data: dict) -> None:
    """Persist a record to disk as JSON.

    Args:
        record_id: Unique identifier for the record.
        data: All data fields to save, including images list.
    """
    import json
    filepath = _record_filepath(record_id)
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2)


def load_record(record_id: str) -> dict | None:
    """Load a record from disk if it exists.

    Args:
        record_id: Unique identifier.

    Returns:
        The loaded dictionary or None if the file does not exist.
    """
    import json
    filepath = _record_filepath(record_id)
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    return None


def list_records() -> list[dict]:
    """Return a list of minimal information about all saved records.

    Each record is represented as a dictionary with keys 'id',
    'company_name' and 'se_name'. Additional keys may be present in
    future but are not used by the index page.
    """
    import json
    records = []
    for filename in os.listdir(app.config['DATA_FOLDER']):
        if not filename.endswith('.json'):
            continue
        record_id = filename[:-5]
        filepath = os.path.join(app.config['DATA_FOLDER'], filename)
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
            records.append({
                'id': record_id,
                'company_name': data.get('company_name', ''),
                'se_name': data.get('se_name', ''),
                'survey_date': data.get('survey_date', ''),
            })
        except Exception:
            continue
    return records


def load_se_names() -> list[str]:
    """Load the list of Solution Engineer names from disk."""
    import json
    try:
        with open(SE_NAMES_FILE, 'r', encoding='utf-8') as f:
            names = json.load(f)
            if isinstance(names, list):
                return names
    except Exception:
        pass
    return []


def save_se_names(names: list[str]) -> None:
    """Save the list of Solution Engineer names to disk."""
    import json
    with open(SE_NAMES_FILE, 'w', encoding='utf-8') as f:
        json.dump(names, f, indent=2)


# HTML templates for the list page and record page.
INDEX_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>Site Survey Reports</title>
  <style>
    body { margin: 0; font-family: Arial, sans-serif; background-color: #0F2C45; color: #FFFFFF; }
    header { display: flex; align-items: center; padding: 1rem 2rem; background-color: #001F3F; }
    header img { height: 36px; margin-right: 1rem; }
    header h1 { font-size: 1.5rem; margin: 0; }
    .container { padding: 2rem; }
    .toolbar { display: flex; justify-content: space-between; align-items: center; margin-bottom: 1rem; }
    .toolbar input[type="text"] { padding: 0.5rem; width: 50%; border-radius: 4px; border: none; }
    .toolbar a { padding: 0.5rem 1rem; background-color: #00B3C7; color: #0F2C45; text-decoration: none; border-radius: 4px; font-weight: bold; }
    table { width: 100%; border-collapse: collapse; margin-top: 1rem; }
    th, td { padding: 0.75rem; text-align: left; border-bottom: 1px solid #324A6D; }
    th { background-color: #001F3F; }
    tr:hover { background-color: #183A59; }
    a.record-link { color: #00B3C7; text-decoration: none; }
  </style>
</head>
<body>
  <header>
    <img src="{{ url_for('static', filename=logo_filename) }}" alt="Cloud Inventory Logo">
    <h1>Site Survey Reports</h1>
  </header>
  <div class="container">
    <div class="toolbar">
      <form method="get" style="width: 100%;">
        <input type="text" name="q" placeholder="Search by company or SE name" value="{{ search_query }}">
      </form>
      <a href="{{ url_for('record_page', record_id='new') }}">Add New</a>
    </div>
    <table>
      <thead>
        <tr>
          <th>Company</th>
          <th>SE Name</th>
          <th>Survey Date</th>
        </tr>
      </thead>
      <tbody>
        {% for rec in records %}
        <tr>
          <td><a class="record-link" href="{{ url_for('record_page', record_id=rec.id) }}">{{ rec.company_name or '(No Name)' }}</a></td>
          <td>{{ rec.se_name or '-' }}</td>
          <td>{{ rec.survey_date or '-' }}</td>
        </tr>
        {% else %}
        <tr>
          <td colspan="3">No records found.</td>
        </tr>
        {% endfor %}
      </tbody>
    </table>
  </div>
</body>
</html>
"""


RECORD_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>{{ 'New' if record_id == 'new' else 'Edit' }} Site Survey Report</title>
  <style>
    body { margin: 0; font-family: Arial, sans-serif; background-color: #0F2C45; color: #FFFFFF; }
    header { display: flex; align-items: center; padding: 1rem 2rem; background-color: #001F3F; }
    header img { height: 36px; margin-right: 1rem; }
    header h1 { font-size: 1.5rem; margin: 0; }
    .container { padding: 2rem; max-width: 900px; margin: 0 auto; }
    label { font-weight: bold; display: block; margin-top: 1rem; }
    input[type=text], textarea, select { width: 100%; padding: 0.5rem; margin-top: 0.25rem; border-radius: 4px; border: none; }
    textarea { resize: vertical; }
    input[type=file] { margin-top: 0.5rem; }
    button.primary { margin-top: 1.5rem; padding: 0.75rem 1.5rem; background-color: #00B3C7; color: #0F2C45; border: none; border-radius: 4px; font-weight: bold; cursor: pointer; }
    button.add-row { margin-top: 0.5rem; padding: 0.5rem 1rem; background-color: #324A6D; color: #FFFFFF; border: none; border-radius: 4px; cursor: pointer; font-size: 0.9rem; }
    table.reps-table { width: 100%; border-collapse: collapse; margin-top: 0.5rem; }
    table.reps-table td { padding: 0.25rem; }
    h3 { margin-top: 1.5rem; margin-bottom: 0.5rem; color: #00B3C7; }
    a.back-link { display: inline-block; margin-bottom: 1rem; color: #00B3C7; text-decoration: none; }
    .operation-section { margin-top: 1.5rem; padding: 1rem; background-color: #122D47; border-radius: 6px; }
    .operation-section h4 { margin-top: 0; color: #00B3C7; }
  </style>
  <script>
    // Add a new row to a representatives table
    function addRow(tableId, nameField, titleField) {
      const table = document.getElementById(tableId);
      const row = document.createElement('tr');
      const nameTd = document.createElement('td');
      const titleTd = document.createElement('td');
      const nameInput = document.createElement('input');
      nameInput.setAttribute('type', 'text');
      nameInput.setAttribute('name', nameField);
      nameInput.style.width = '100%';
      const titleInput = document.createElement('input');
      titleInput.setAttribute('type', 'text');
      titleInput.setAttribute('name', titleField);
      titleInput.style.width = '100%';
      nameTd.appendChild(nameInput);
      titleTd.appendChild(titleInput);
      row.appendChild(nameTd);
      row.appendChild(titleTd);
      table.appendChild(row);
    }

    // Add a custom operations section under the General category
    function addCustomSection() {
      const sectionName = prompt('Enter the name of the new operations section:');
      if (!sectionName) return;
      // Sanitize section name to use in field names
      const key = sectionName.replace(/\s+/g, '_');
      // Update hidden input that tracks dynamic sections
      const hiddenInput = document.getElementById('dynamic_sections');
      let existing = hiddenInput.value ? hiddenInput.value.split(',') : [];
      if (!existing.includes(key)) {
        existing.push(key);
        hiddenInput.value = existing.join(',');
      }
      // Create a new operation section div
      const container = document.createElement('div');
      container.className = 'operation-section';
      const title = document.createElement('h4');
      title.textContent = sectionName;
      container.appendChild(title);
      // Observations
      const obsLabel = document.createElement('label');
      obsLabel.textContent = 'Observations:';
      const obsText = document.createElement('textarea');
      obsText.setAttribute('name', `section_${key}_obs`);
      obsText.rows = 2;
      container.appendChild(obsLabel);
      container.appendChild(obsText);
      // Pain point
      const painLabel = document.createElement('label');
      painLabel.textContent = 'Pain point:';
      const painText = document.createElement('textarea');
      painText.setAttribute('name', `section_${key}_pain`);
      painText.rows = 2;
      container.appendChild(painLabel);
      container.appendChild(painText);
      // Root cause
      const rootLabel = document.createElement('label');
      rootLabel.textContent = 'Root cause:';
      const rootText = document.createElement('textarea');
      rootText.setAttribute('name', `section_${key}_root`);
      rootText.rows = 2;
      container.appendChild(rootLabel);
      container.appendChild(rootText);
      // Impact
      const impactLabel = document.createElement('label');
      impactLabel.textContent = 'Business impact:';
      const impactText = document.createElement('textarea');
      impactText.setAttribute('name', `section_${key}_impact`);
      impactText.rows = 2;
      container.appendChild(impactLabel);
      container.appendChild(impactText);
      // Photo upload
      const fileLabel = document.createElement('label');
      fileLabel.textContent = 'Upload photo/file:';
      const fileInput = document.createElement('input');
      fileInput.setAttribute('type', 'file');
      fileInput.setAttribute('name', `section_${key}_photo`);
      fileInput.setAttribute('accept', 'image/*');
      container.appendChild(fileLabel);
      container.appendChild(fileInput);
      // Append new section after the existing operations container
      document.getElementById('operations-container').appendChild(container);
    }
  </script>
</head>
<body>
  <header>
    <img src="{{ url_for('static', filename=logo_filename) }}" alt="Cloud Inventory Logo">
    <h1>{{ 'New' if record_id == 'new' else 'Edit' }} Site Survey Report</h1>
  </header>
  <div class="container">
    <a class="back-link" href="{{ url_for('index_page') }}">&larr; Back to list</a>
    <form method="post" enctype="multipart/form-data">
      <!-- Solution Engineer selection -->
      <label>Solution Engineer:</label>
      <select name="se_name_select">
        <option value="">-- Select --</option>
        {% for name in se_names %}
        <option value="{{ name }}" {% if name == data.get('se_name') %}selected{% endif %}>{{ name }}</option>
        {% endfor %}
      </select>
      <label>Add new Solution Engineer:</label>
      <input type="text" name="new_se_name" placeholder="Enter new SE name">
      <!-- Basic fields -->
      <label>Survey location:</label>
      <input type="text" name="survey_location" required value="{{ data.get('survey_location','') }}">
      <label>Survey conducted on:</label>
      <input type="text" name="survey_date" required value="{{ data.get('survey_date','') }}" placeholder="YYYY-MM-DD">
      <label>Company name:</label>
      <input type="text" name="company_name" required value="{{ data.get('company_name','') }}">
      <label>Company & Site Profile:</label>
      <textarea name="company_profile" rows="3">{{ data.get('company_profile','') }}</textarea>
      <label>Upload photo/file for Company & Site Profile:</label>
      <input type="file" name="company_profile_file" accept="image/*">
      {% if data.get('company_profile_image') %}
      <div class="images-list">
        <img src="{{ url_for('static', filename=data.get('company_profile_image')) }}" alt="Company profile" style="max-width:200px; margin-top:0.5rem;">
      </div>
      {% endif %}
      <label>Master Data:</label>
      <textarea name="master_data_text" rows="3">{{ data.get('master_data_text','') }}</textarea>
      <label>Upload file/photo for Master Data:</label>
      <input type="file" name="master_data_file" accept="image/*">
      {% if data.get('master_data_image') %}
      <div class="images-list">
        <img src="{{ url_for('static', filename=data.get('master_data_image')) }}" alt="Master Data" style="max-width:200px; margin-top:0.5rem;">
      </div>
      {% endif %}
      <label>Inventory Information:</label>
      <textarea name="inventory_text" rows="3">{{ data.get('inventory_text','') }}</textarea>
      <label>Upload file/photo for Inventory Information:</label>
      <input type="file" name="inventory_file" accept="image/*">
      {% if data.get('inventory_image') %}
      <div class="images-list">
        <img src="{{ url_for('static', filename=data.get('inventory_image')) }}" alt="Inventory Info" style="max-width:200px; margin-top:0.5rem;">
      </div>
      {% endif %}
      <label>IT & Systems Landscape:</label>
      <textarea name="it_systems" rows="3">{{ data.get('it_systems','') }}</textarea>
      <!-- Representatives -->
      <h3>Company Representatives</h3>
      <!-- Labels indicating fields -->
      <div style="display:flex; gap:2rem; margin-top:0.25rem;">
        <span style="color:#00B3C7; font-weight:bold;">Name</span>
        <span style="color:#00B3C7; font-weight:bold;">Position</span>
      </div>
      <table id="company-reps" class="reps-table">
        <tr><th style="text-align:left; color:#00B3C7;">Name</th><th style="text-align:left; color:#00B3C7;">Position</th></tr>
        {% for rep in data.get('company_representatives', []) %}
        <tr>
          <td><input type="text" name="company_rep_name[]" value="{{ rep.name }}" style="width:100%"></td>
          <td><input type="text" name="company_rep_title[]" value="{{ rep.title }}" style="width:100%"></td>
        </tr>
        {% endfor %}
        <!-- Always include one blank row -->
        <tr>
          <td><input type="text" name="company_rep_name[]" value="" style="width:100%"></td>
          <td><input type="text" name="company_rep_title[]" value="" style="width:100%"></td>
        </tr>
      </table>
      <button type="button" class="add-row" onclick="addRow('company-reps','company_rep_name[]','company_rep_title[]')">Add Representative</button>
      <h3>Cloud Inventory Representatives</h3>
      <!-- Labels indicating fields -->
      <div style="display:flex; gap:2rem; margin-top:0.25rem;">
        <span style="color:#00B3C7; font-weight:bold;">Name</span>
        <span style="color:#00B3C7; font-weight:bold;">Position</span>
      </div>
      <table id="ci-reps" class="reps-table">
        <tr><th style="text-align:left; color:#00B3C7;">Name</th><th style="text-align:left; color:#00B3C7;">Position</th></tr>
        {% for rep in data.get('ci_representatives', []) %}
        <tr>
          <td><input type="text" name="ci_rep_name[]" value="{{ rep.name }}" style="width:100%"></td>
          <td><input type="text" name="ci_rep_title[]" value="{{ rep.title }}" style="width:100%"></td>
        </tr>
        {% endfor %}
        <tr>
          <td><input type="text" name="ci_rep_name[]" value="" style="width:100%"></td>
          <td><input type="text" name="ci_rep_title[]" value="" style="width:100%"></td>
        </tr>
      </table>
      <button type="button" class="add-row" onclick="addRow('ci-reps','ci_rep_name[]','ci_rep_title[]')">Add CI Representative</button>
      <!-- Operations sections -->
      <h3>Operations Details</h3>
      <input type="hidden" name="dynamic_sections" id="dynamic_sections" value="{{ dynamic_sections }}">
      <div id="operations-container">
      {% for key in operations_keys %}
      <div class="operation-section">
        <h4>{{ key }}</h4>
        <label>Observations:</label>
        <textarea name="section_{{ key }}_obs" rows="2">{{ data.get('operations', {}).get(key, {}).get('observations','') }}</textarea>
        <label>Pain point:</label>
        <textarea name="section_{{ key }}_pain" rows="2">{{ data.get('operations', {}).get(key, {}).get('pain','') }}</textarea>
        <label>Root cause:</label>
        <textarea name="section_{{ key }}_root" rows="2">{{ data.get('operations', {}).get(key, {}).get('root_cause','') }}</textarea>
        <label>Business impact:</label>
        <textarea name="section_{{ key }}_impact" rows="2">{{ data.get('operations', {}).get(key, {}).get('impact','') }}</textarea>
        <label>Upload photo/file:</label>
        <input type="file" name="section_{{ key }}_photo" accept="image/*">
        {% if data.get('operations', {}).get(key, {}).get('images') %}
        <div class="images-list">
          {% for img in data.get('operations')[key]['images'] %}
          <img src="{{ url_for('static', filename=img) }}" alt="Uploaded">
          {% endfor %}
        </div>
        {% endif %}
      </div>
      {% endfor %}
      </div>
      <!-- Add custom operations section button -->
      <button type="button" class="add-row" onclick="addCustomSection()">Add New Section</button>
      <!-- Additional attachments -->
      <label>Other photos (optional):</label>
      <input type="file" name="photos" multiple accept="image/*">
      {% if images %}
      <div class="images-list">
        {% for img in images %}
        <img src="{{ url_for('static', filename=img) }}" alt="Uploaded image">
        {% endfor %}
      </div>
      {% endif %}
      <!-- In-scope / out-of-scope -->
      <label>In‑Scope Requirements:</label>
      <textarea name="in_scope" rows="2">{{ data.get('in_scope','') }}</textarea>
      <label>Out‑of‑Scope Requirements:</label>
      <textarea name="out_scope" rows="2">{{ data.get('out_scope','') }}</textarea>
      <label>Value Summary & Metrics (optional):</label>
      <textarea name="value_summary" rows="3" placeholder="Baseline metrics, target improvements, ROI assumptions.">{{ data.get('value_summary','') }}</textarea>
      <button type="submit" class="primary">Save & Download Report</button>
    </form>
  </div>
</body>
</html>
"""


@app.route('/new')
def new_redirect():
    """Redirect the legacy /new path to the new record page."""
    return redirect(url_for('record_page', record_id='new'))


@app.route('/')
def index_page():
    """Display a list of existing records with search and an option to create new."""
    # Choose the negative logo for dark background
    logo_filename = 'ci-negative.png'
    # Extract search query from query parameters
    search_query = request.args.get('q', '').strip().lower()
    records = list_records()
    # Filter records based on search query (match company or SE name)
    if search_query:
        records = [r for r in records if search_query in (r.get('company_name','').lower() + r.get('se_name','').lower())]
    # Sort records by survey_date descending
    records.sort(key=lambda r: r.get('survey_date', ''), reverse=True)
    return render_template_string(INDEX_TEMPLATE, records=records, logo_filename=logo_filename, search_query=search_query)

@app.route('/record/<record_id>', methods=['GET', 'POST'])
def record_page(record_id: str):
    """Create or edit a site survey record.

    Args:
        record_id: 'new' for a new record or an existing record ID.
    """
    logo_filename = 'ci-negative.png'
    if request.method == 'POST':
        # Determine record ID if new
        if record_id == 'new':
            record_id = uuid.uuid4().hex
        # Load existing record if editing
        existing = load_record(record_id)
        saved_images = existing.get('images', []) if existing else []
        # Solution Engineer selection
        selected_se = request.form.get('se_name_select', '').strip()
        new_se = request.form.get('new_se_name', '').strip()
        se_names = load_se_names()
        if new_se:
            # Use the new SE name and persist it
            se_name = new_se
            if se_name not in se_names:
                se_names.append(se_name)
                save_se_names(se_names)
        else:
            se_name = selected_se
        # Gather form inputs
        data = {
            'se_name': se_name,
            'survey_location': request.form.get('survey_location','').strip(),
            'survey_date': request.form.get('survey_date','').strip(),
            'company_name': request.form.get('company_name','').strip(),
            'company_profile': request.form.get('company_profile','').strip(),
            # separate master data and inventory text fields
            'master_data_text': request.form.get('master_data_text','').strip(),
            'inventory_text': request.form.get('inventory_text','').strip(),
            'it_systems': request.form.get('it_systems','').strip(),
            # Observations fields below are kept for backward compatibility
            'observations': request.form.get('observations','').strip(),
            'root_cause': request.form.get('root_cause','').strip(),
            'business_impact': request.form.get('business_impact','').strip(),
            'in_scope': request.form.get('in_scope','').strip(),
            'out_scope': request.form.get('out_scope','').strip(),
            'value_summary': request.form.get('value_summary','').strip(),
        }
        # Company and CI representatives (multi-row)
        comp_names = request.form.getlist('company_rep_name[]')
        comp_titles = request.form.getlist('company_rep_title[]')
        ci_names = request.form.getlist('ci_rep_name[]')
        ci_titles = request.form.getlist('ci_rep_title[]')
        company_reps = []
        for name, title in zip(comp_names, comp_titles):
            if name.strip() or title.strip():
                company_reps.append({'name': name.strip(), 'title': title.strip()})
        ci_reps = []
        for name, title in zip(ci_names, ci_titles):
            if name.strip() or title.strip():
                ci_reps.append({'name': name.strip(), 'title': title.strip()})
        data['company_representatives'] = company_reps
        data['ci_representatives'] = ci_reps
        # Operations details
        static_ops = ['Receiving', 'Putaway', 'Replenishment', 'Order Management', 'Picking', 'Packing', 'Staging', 'Shipping', 'General']
        # Dynamic sections specified by the client
        dynamic_sections_str = request.form.get('dynamic_sections', '').strip()
        dynamic_ops = [s for s in dynamic_sections_str.split(',') if s]
        operations_keys = static_ops + dynamic_ops
        operations_data = {}
        # Ensure a static directory exists for this record for operations images
        record_static_dir = os.path.join(app.static_folder, record_id)
        os.makedirs(record_static_dir, exist_ok=True)
        for key in operations_keys:
            obs = request.form.get(f'section_{key}_obs', '').strip()
            pain = request.form.get(f'section_{key}_pain', '').strip()
            root = request.form.get(f'section_{key}_root', '').strip()
            impact = request.form.get(f'section_{key}_impact', '').strip()
            op_images = []
            # Each operations file input returns one file
            file = request.files.get(f'section_{key}_photo')
            if file and file.filename:
                ext = os.path.splitext(file.filename)[1]
                unique_name = f"{uuid.uuid4().hex}{ext}"
                img_path = os.path.join(record_static_dir, unique_name)
                file.save(img_path)
                rel_path = f"{record_id}/{unique_name}"
                op_images.append(rel_path)
                saved_images.append(rel_path)
            operations_data[key] = {
                'observations': obs,
                'pain': pain,
                'root_cause': root,
                'impact': impact,
                'images': op_images,
            }
        data['operations'] = operations_data

        # Handle single uploads for company/site profile, master data and inventory sections
        # Save images to record_static_dir and store relative paths in data
        data['company_profile_image'] = ''
        cp_file = request.files.get('company_profile_file')
        if cp_file and cp_file.filename:
            ext = os.path.splitext(cp_file.filename)[1]
            unique_name = f"{uuid.uuid4().hex}{ext}"
            img_path = os.path.join(record_static_dir, unique_name)
            cp_file.save(img_path)
            rel_path = f"{record_id}/{unique_name}"
            data['company_profile_image'] = rel_path
        data['master_data_image'] = ''
        md_file = request.files.get('master_data_file')
        if md_file and md_file.filename:
            ext = os.path.splitext(md_file.filename)[1]
            unique_name = f"{uuid.uuid4().hex}{ext}"
            img_path = os.path.join(record_static_dir, unique_name)
            md_file.save(img_path)
            rel_path = f"{record_id}/{unique_name}"
            data['master_data_image'] = rel_path
        data['inventory_image'] = ''
        inv_file = request.files.get('inventory_file')
        if inv_file and inv_file.filename:
            ext = os.path.splitext(inv_file.filename)[1]
            unique_name = f"{uuid.uuid4().hex}{ext}"
            img_path = os.path.join(record_static_dir, unique_name)
            inv_file.save(img_path)
            rel_path = f"{record_id}/{unique_name}"
            data['inventory_image'] = rel_path
        # Handle general file uploads (photos) that are not tied to operations
        photos = request.files.getlist('photos') or []
        for photo in photos:
            if photo.filename:
                ext = os.path.splitext(photo.filename)[1]
                unique_name = f"{uuid.uuid4().hex}{ext}"
                img_path = os.path.join(record_static_dir, unique_name)
                photo.save(img_path)
                rel_path = f"{record_id}/{unique_name}"
                saved_images.append(rel_path)
        # Generate summary for MCP agent based on operations data and company profile
        summary_parts = [data['company_profile']]
        for op_key, op_val in operations_data.items():
            if op_val['pain']:
                summary_parts.append(f"{op_key} pain: {op_val['pain']}")
            if op_val['impact']:
                summary_parts.append(f"Impact: {op_val['impact']}")
        summary = '\n'.join([p for p in summary_parts if p])
        mcp_response = call_mcp_agent(summary)
        # Save record data including images list
        data_for_save = data.copy()
        data_for_save['images'] = saved_images
        save_record(record_id, data_for_save)
        # Build absolute paths for all images (general and operations) for Word report
        abs_paths = [os.path.join(app.static_folder, rel) for rel in saved_images]
        report_path = generate_word_report(data, abs_paths, mcp_response)
        return send_file(report_path, as_attachment=True)
    else:
        # GET request: load record or create blank for new
        record_data = load_record(record_id) if record_id != 'new' else None
        if not record_data:
            # Initialize blank record data structure
            record_data = {
                'se_name': '', 'survey_location': '', 'survey_date': '', 'company_name': '',
                'company_profile': '', 'master_data_text': '', 'inventory_text': '', 'it_systems': '',
                'company_profile_image': '', 'master_data_image': '', 'inventory_image': '',
                'observations': '', 'root_cause': '', 'business_impact': '',
                'in_scope': '', 'out_scope': '', 'value_summary': '',
                'company_representatives': [], 'ci_representatives': [],
                'operations': {k: {'observations': '', 'pain': '', 'root_cause': '', 'impact': '', 'images': []}
                              for k in ['Receiving','Putaway','Replenishment','Order Management','Picking','Packing','Staging','Shipping','General']},
                'images': [],
            }
        # Extract list of images (general and operations images) for display
        images = record_data.get('images', [])
        # Provide list of SE names
        se_names = load_se_names()
        # Determine operations keys from record_data to include dynamic sections
        operations_keys = list(record_data.get('operations', {}).keys())
        # Compute dynamic sections string for hidden input (exclude static ones)
        static_ops = ['Receiving','Putaway','Replenishment','Order Management','Picking','Packing','Staging','Shipping','General']
        dynamic_sections = ','.join([k for k in operations_keys if k not in static_ops])
        return render_template_string(
            RECORD_TEMPLATE,
            record_id=record_id,
            data=record_data,
            images=images,
            logo_filename=logo_filename,
            se_names=se_names,
            operations_keys=operations_keys,
            dynamic_sections=dynamic_sections
        )


if __name__ == '__main__':
    # For local development.  Render runs the app using Gunicorn.
    app.run(debug=True)